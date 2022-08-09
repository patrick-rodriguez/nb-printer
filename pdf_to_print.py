from docx import *
from tkinter import *
from PyPDF2 import PdfWriter, PdfReader
import win32api
import win32print
from docx2pdf import convert
from docx.shared import Mm
from docx.shared import Inches
from docx.shared import Pt
import pandas as pd

#----------Convert excel file to dataframe-------#
spreadsheet = pd.read_excel("descriptions.xlsx") #
#------------------------------------------------#

#-Convert and sort list of rocks for user interface-#
rock_name_list = spreadsheet["Rock Name"].to_list() #
sorted_rock_name_list = rock_name_list.copy()       #
sorted_rock_name_list.sort()                        #
#---------------------------------------------------#

#-----------------------Creating the user interface--------------------------------------------#
window = Tk()                                                                                  #
window.geometry("1000x1000")                                                                   #
                                                                                               #
main_frame = Frame(window)                                                                     #
main_frame.pack(fill=BOTH, expand=1)                                                           #
                                                                                               #
canvas = Canvas(main_frame)                                                                    #
canvas.pack(side=LEFT, fill=BOTH, expand=1)                                                    #
                                                                                               #
scrollbar = Scrollbar(canvas, orient="horizontal", command=canvas.xview)                   #
scrollbar.pack(side=BOTTOM, fill="x")                                                            #
                                                                                               #
canvas.configure(xscrollcommand=scrollbar.set)                                                 #
canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))        #
                                                                                               #
second_frame = Frame(canvas)                                                                   #
canvas.create_window((0, 0), window=second_frame, anchor="nw")                                 #
#----------------------------------------------------------------------------------------------#

#-------------------------Generate PDF of formatted microsoft word document------------------------#
def create_pdf(button_text):                                                                       #
    document = Document()                                                                          #
    section = document.sections[0]                                                                 #
    section.page_height = Mm(297)                                                                  #
    section.page_width = Mm(210)                                                                   #
    section.left_margin = Mm(10.4)  # originally 25.4                                              #
    section.right_margin = Mm(10.4)                                                                #
    section.top_margin = Mm(10.4)  # originally 25.4                                               #
    section.bottom_margin = Mm(10.4)                                                               #
    section.header_distance = Mm(12.7)                                                             #
    section.footer_distance = Mm(12.7)                                                             #
    table = document.add_table(rows=1, cols=2, style='Table Grid')                                 #
#--------------------------------------------------------------------------------------------------#

#---Grab the respective cells in the excel sheet for the rock name and description---#
    excel_name_cell = ""                                                             #
    excel_description_cell = ""                                                      #
    for i in range(len(rock_name_list)):                                             #
        if button_text == rock_name_list[i]:                                         #
            excel_name_cell = spreadsheet.iat[i, 0]                                  #
            excel_description_cell = spreadsheet.iat[i, 1]                           #
#------------------------------------------------------------------------------------#

#---Create the table and fill the cells with the rock name and description---#
    for new_row in range(30):                                                #
        table.add_row()                                                      #
                                                                             #
    for row in table.rows:                                                   #
        for cell in row.cells:                                               #
            cell.width = Inches(8)                                           #
            para = cell.paragraphs[0]                                        #
            run = para.add_run(excel_name_cell)                              #
            run.bold = True                                                  #
            run.underline = True                                             #
            run.font.name = "Arial"                                          #
            run.font.size = Pt(8)                                            #
                                                                             #
            run = para.add_run(excel_description_cell + "\n")                #
            run.font.name = "Arial"                                          #
            run.font.size = Pt(8)                                            #
                                                                             #
    document.save("test.docx")                                               #
    convert("test.docx", r"C:\Users\patri\OneDrive\Documents\test.pdf")                                         #

    infile = PdfReader(r"C:\Users\patri\OneDrive\Documents\test.pdf")
    output = PdfWriter()

    output.add_page(infile.pages[0])

    with open(r"C:\Users\patri\OneDrive\Documents\test.pdf", "wb") as f:
        output.write(f)

    win32api.ShellExecute(0,
                          "print",
                         r"C:\Users\patri\OneDrive\Documents\test.pdf",
                          'd: "%s"' % win32print.GetDefaultPrinter(),
                          ".",
                          0
                          )
#----------------------------------------------------------------------------#

number_of_rocks = len(rock_name_list)
row_number = 0
column_number = 0
for i in range(number_of_rocks):
    rock_name_button = Button(second_frame, text=sorted_rock_name_list[i], width=30, command=lambda button_text=sorted_rock_name_list[i]: create_pdf(button_text))
    if i % 30 == 0:
        row_number = 0
        column_number += 1
    rock_name_button.grid(row=row_number, column=column_number)
    row_number += 1

window.mainloop()

