from tkinter import *
from customtkinter import *
import pandas as pd
from docx import *
from PyPDF2 import PdfWriter, PdfReader
import win32api
import win32print
from docx2pdf import convert
from docx.shared import Mm
from docx.shared import Inches
from docx.shared import Pt

window = Tk()
window.state("zoomed")
window.title("Nature's Bling Description Printer")
window.configure(background="grey21")

wrapper1 = CTkFrame(window, height=400, fg_color="grey", bg_color="grey")
wrapper2 = CTkFrame(window, width=100, fg_color="black", bg_color="black", border_color="black")

wrapper1.pack_propagate(0)
wrapper2.pack_propagate(0)

wrapper1.pack(fill="both", expand="yes", padx=10, pady=10)
wrapper2.pack(fill="both", expand="yes", padx=10, pady=10)

canvas = CTkCanvas(wrapper1, bg="grey21")
canvas.pack_propagate(0)
canvas.pack(side="left", fill="both", expand=1)

canvas2 = CTkCanvas(wrapper2, bg="grey", bd=5)
canvas2.pack(side="left", fill="both", expand=1)

wrapper1_scrollbar = Scrollbar(canvas, orient="horizontal", command=canvas.xview)
wrapper1_scrollbar.pack(side=BOTTOM, fill="x")
canvas.configure(xscrollcommand=wrapper1_scrollbar.set)  #
canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

button_frame = Frame(canvas, bg="grey21")
canvas.create_window((0, 0), window=button_frame, anchor="nw")

spreadsheet = pd.read_excel("descriptions.xlsx")
rock_name_list = spreadsheet["Rock Name"].to_list()
sorted_rock_name_list = rock_name_list.copy()
sorted_rock_name_list.sort()

checkbox_list = []
queue = []


def add_to_queue(text, queue):
    if len(queue) <= 10:
        already_in_queue = False
        for box in queue:
            if box == text:
                already_in_queue = True
                break
        if not already_in_queue:
            checkbox = CTkCheckBox(canvas2, text=text, text_font=("Trajan Pro", 15), text_color="white")
            checkbox.configure(command=lambda: [checkbox.destroy(), queue.remove(text)])
            checkbox.pack(anchor="nw")
            queue.append(checkbox.text)
            checkbox_list.append(checkbox)


def clear():
    for i in range(len(queue)):
        queue.remove(queue[0])

    for j in range(len(checkbox_list)):
        checkbox_list[0].destroy()
        checkbox_list.remove(checkbox_list[0])


def generate_pdf(button_text):
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(10.4)  # originally 25.4
    section.right_margin = Mm(10.4)
    section.top_margin = Mm(10.4)  # originally 25.4
    section.bottom_margin = Mm(10.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    table = document.add_table(rows=1, cols=2, style='Table Grid')

    excel_name_cell = ""
    excel_description_cell = ""
    for i in range(len(rock_name_list)):
        if button_text == rock_name_list[i]:
            excel_name_cell = spreadsheet.iat[i, 0]
            excel_description_cell = spreadsheet.iat[i, 1]

    for new_row in range(30):
        table.add_row()

    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(8)
            para = cell.paragraphs[0]
            run = para.add_run(excel_name_cell)
            run.bold = True
            run.underline = True
            run.font.name = "Arial"
            run.font.size = Pt(8)

            run = para.add_run(excel_description_cell + "\n")
            run.font.name = "Arial"
            run.font.size = Pt(8)

    document.save("test.docx")
    convert("test.docx", r"C:\Users\patri\OneDrive\Documents\test.pdf")  #

    infile = PdfReader(r"C:\Users\patri\OneDrive\Documents\test.pdf")
    output = PdfWriter()

    output.add_page(infile.pages[0])

    with open(r"C:\Users\patri\OneDrive\Documents\test.pdf", "wb") as f:
        output.write(f)

    win32api.ShellExecute(0,
                          "print",
                          r"C:\Users\patri\OneDrive\Documents\test.pdf",
                          'd: "%s"' % win32print.GetDefaultPrinter(),
                          ".",
                          0
                          )

def print_queue():
    for i in range(len(queue)):
        generate_pdf(queue[i])
        clear()


print_button = CTkButton(canvas2, text="PRINT", width=402, height=296, command=print_queue)
print_button.place(x=705, y=2)

clear_button = CTkButton(canvas2, text="CLEAR", width=402, height=296, command=clear)
clear_button.place(x=1110, y=2)

number_of_rocks = len(rock_name_list)
row_number = 0
column_number = 0
for i in range(number_of_rocks):
    rock_name_button = CTkButton(button_frame,
                                 text_font=("Times", 10),
                                 padx=2,
                                 pady=2.4,
                                 width=300,
                                 text=sorted_rock_name_list[i],
                                 command=lambda button_text=sorted_rock_name_list[i]: add_to_queue(button_text, queue))
    if i % 15 == 0:
        row_number = 0
        column_number += 1
    rock_name_button.grid(row=row_number, column=column_number)
    row_number += 1

window.mainloop()
